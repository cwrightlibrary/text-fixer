from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from os.path import dirname, join, realpath
import re

current_dir = dirname(realpath(__file__))

class Jobs:
    def __init__(self, input_doc):
        self.doc_list = []
        self.orig_doc = Document(join(current_dir, input_doc))
    
    def get_job_info(self):
        for p in self.orig_doc.paragraphs:
            if "submission date" in p.text.lower():
                _date = p.text[17:]
            elif p.text[0:11] == "Job Title: ":
                _title = p.text[11:]
            elif p.text[0:17] == "Job Description: ":
                _desc = p.text[17:]
            elif p.text[0:16] == "Qualifications: ":
                _qual = p.text[16:]
            elif p.text[0:6] == "Type: ":
                _type = p.text[0:6]
            elif p.text[0:14] == "How to Apply: ":
                _how = p.text[14:]
            elif p.text[0:7] == "Salary: ":
                _salary = p.text[7:]
            elif p.text[0:9] == "Contact: ":
                _contact = p.text[9:]
            elif p.text[0:7] == "!BREAK!":
                job = {
                    "date": _date,
                    "title": _title,
                    "desc": _desc,
                    "qual": _qual,
                    "type": _type,
                    "how": _how,
                    "salary": _salary,
                    "contact": _contact
                }
                self.doc_list.append(job)
            _date, _title, _desc, _qual, _type, _how, _salary, _contact = "", "", "", "", "", "", "", ""
    
    def fix_element(self):
        temp_list = self.doc_list
        self.doc_list = []
        for doc in temp_list:
            new_doc = {}
            new_doc["Submission Date: "] = doc["date"]
            new_doc["Job Title: "] = doc["title"]
            new_doc["Job Description: "] = doc["desc"]
            new_doc["Qualifications: "] = doc["qual"]
            new_doc["Type: "] = doc["type"]
            new_doc["How to Apply: "] = doc["type"]
            new_doc["Salary: "] = doc["salary"]
            new_doc["Contact: "] = doc["contact"]
            
            self.doc_list.append(new_doc)
            new_doc = {}
    
    def create_documents(self):
        self.export_doc_list = []
        
        for doc in self.doc_list:
            new_doc = Document()
            new_doc.styles["Normal"].font.name = "Arial"
            new_doc.styles["Heading 2"].font.name = "Arial"
            new_doc.styles["Heading 2"].font.size = Pt(16)
            new_doc.styles["Heading 2"].font.color.rgb = RGBColor(0, 0, 0)
            
            h = new_doc.add_heading("", 2)
            h.add_run(doc["Job Title: "] + ", " + doc["Submission Date: "]).bold = False
            
            for k, v in doc.items():
                p = new_doc.add_paragraph()
                p.style = new_doc.styles["Normal"]
                p.add_run(k).bold = True
                p.add_run(v)
                p.add_run()
            
            for p in new_doc.paragraphs:
                if p.style.name.startswith("Heading 2"):
                    for run in p.runs:
                        run.font.name = "Arial"
            
            self.export_doc_list.append(new_doc)
    
    def save_documents(self):
        pass
