from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from os.path import dirname, join, realpath
import re

current_dir = dirname(realpath(__file__))


def get_job_info(doc):
    doc_list = []
    original_doc = Document(join(current_dir, doc))

    for paragraph in original_doc.paragraphs:
        if "submission date" in paragraph.text.lower():
            job_date = paragraph.text[17:]
        if paragraph.text[0:17] == "Submission Date: ":
            job_date = paragraph.text[17:]
        elif paragraph.text[0:11] == "Job Title: ":
            job_title = paragraph.text[11:]
        elif paragraph.text[0:17] == "Job Description: ":
            job_desc = paragraph.text[17:]
        elif paragraph.text[0:16] == "Qualifications: ":
            job_qual = paragraph.text[16:]
        elif paragraph.text[0:6] == "Type: ":
            job_type = paragraph.text[6:]
        elif paragraph.text[0:14] == "How to Apply: ":
            job_how = paragraph.text[14:]
        elif paragraph.text[0:7] == "Salary:":
            job_salary = paragraph.text[7:]
        elif paragraph.text[0:9] == "Contact: ":
            job_contact = paragraph.text[9:]
        elif paragraph.text[0:7] == "!BREAK!":
            job_how = (
                job_how.replace("APPLICATION PROCEDURE", "")
                .replace("•", "\n-")
                .replace("Application Instructions", "")
                .replace("  ", " ")
            )

            job = {
                "date": job_date,
                "title": job_title,
                "desc": job_desc,
                "qual": job_qual,
                "type": job_type,
                "how": job_how,
                "salary": job_salary,
                "contact": job_contact,
            }
            doc_list.append(job)
            (
                job_date,
                job_title,
                job_desc,
                job_qual,
                job_type,
                job_how,
                job_salary,
                job_contact,
            ) = ("", "", "", "", "", "", "", "")

    return doc_list


def fix_date(element):
    new_element = element

    return new_element


def fix_title(element):
    new_element = element
    new_element = new_element.replace("and", "&").title()

    return new_element


def fix_description(element):
    new_element = element
    new_element = new_element.replace("eResources", "e!!Resources")
    new_element = new_element.replace("•", "-").replace("●", "-")
    new_element = new_element.replace(")-", ")\n-")
    new_element = new_element.replace("yX", "y X")
    new_element = new_element.replace("--", "")
    new_element = new_element.replace(" | ", "\n").replace("|", "\n")
    new_element = new_element.replace(".$", ". $").replace(".*", ". *")
    new_element = (
        new_element.replace("Description ", "")
        .replace("Essential Job Duties ", "")
        .replace("ADDITIONAL DUTIES ", "")
        .replace("JOB SUMMARY", "")
        .replace("ADDITIONAL DUTIES: ", "")
    )

    letter_spacing = re.findall("[a-z][A-Z]", new_element)
    colon_spacing = re.findall(":[a-zA-Z]", new_element)
    colon_spacing_2 = re.findall(":[1-9]", new_element)
    period_spacing = re.findall(".[A-Z]", new_element)
    dash_spacing = re.findall("[a-z]-[A-Z]", new_element)

    special_spacing = colon_spacing + colon_spacing_2 + period_spacing

    for e in special_spacing:
        if e[0] == ":":
            new_element = new_element.replace(e, e[0] + "\n" + e[1])
        elif e[0] == ".":
            new_element = new_element.replace(e, e[0] + " " + e[1])
        elif e[1] == "-":
            new_element = new_element.replace(e, e[0] + "\n" + e[1])

    for e in letter_spacing:
        new_element = new_element.replace(e, e[0] + ".\n" + e[1])

    for e in dash_spacing:
        new_element = new_element.replace(e, e[0] + "\n" + e[1])

    new_element = new_element.replace("e!!Resources", "eResources")

    if new_element[len(new_element) - 1 :] != ".":
        new_element += "."

    return new_element


def fix_qualifications(element):
    new_element = element
    new_element = new_element.replace("•", "-")
    new_element = new_element.replace(")-", ")\n-")
    new_element = new_element.replace("yX", "y X")
    new_element = new_element.replace("--", "")

    letter_spacing = re.findall("[a-z][A-Z]", new_element)
    colon_spacing = re.findall(":[a-zA-Z]", new_element)
    colon_spacing_2 = re.findall(":[1-9]", new_element)
    period_spacing = re.findall(".[a-zA-Z]", new_element)
    dash_spacing = re.findall("[a-z]-[A-Z]", new_element)

    special_spacing = colon_spacing + colon_spacing_2 + period_spacing

    for e in special_spacing:
        if e[0] == ":":
            new_element = new_element.replace(e, e[0] + "\n" + e[1])
        elif e[0] == ".":
            new_element = new_element.replace(e, e[0] + " " + e[1])

    for e in letter_spacing:
        new_element = new_element.replace(e, e[0] + ".\n" + e[1])

    for e in dash_spacing:
        new_element = new_element.replace(e, e[1] + "\n" + e[2])

    if new_element[len(new_element) - 1 :] != ".":
        new_element += "."

    return new_element


def fix_type(element):
    new_element = element
    if "part" in new_element or "Part" in new_element:
        new_element = "Part-Time"
    elif "full" in new_element or "Full" in new_element:
        new_element = "Full-Time"
    else:
        new_element = new_element.title()

    return new_element


def fix_salary(element):
    new_element = element
    if new_element == "":
        new_element = ""

    return new_element


# TODO add a fix_how() function


def fix_contact(element):
    new_element = element
    new_element = new_element.replace("•", "-")
    new_element = new_element.replace("--", "")
    new_element = new_element.replace(" | ", "\n").replace("|", "\n")
    new_element = (
        new_element.replace(".edu ", ".edu\n")
        .replace(".com ", ".com\n")
        .replace(".net ", ".net\n")
    )

    letter_spacing = re.findall("[a-z][A-Z]", new_element)
    paren_spacing = re.findall("[a-zA-Z]" + r"\(", new_element)

    for e in letter_spacing:
        if e[0] != "h" and e[1] != "D":
            new_element = new_element.replace(e, e[0] + " " + e[1])

    for e in paren_spacing:
        new_element = new_element.replace(e, e[0] + "\n" + e[1])

    phone = re.compile(r"\b(?:\d{3}[-.\s]?)?\d{3}[-.\s]?\d{4}\b").search(new_element)

    if phone:
        phone = phone.group()
        new_element = new_element.replace(" " + phone, ", " + phone)
    else:
        phone = ""

    new_element = new_element.replace("),", ")")

    return new_element


def load_fix_document(loc):
    doc_list = get_job_info(loc)
    new_doc_list = []

    for doc in doc_list:
        new_doc = {}
        new_doc["Submission Date: "] = fix_date(doc["date"])
        new_doc["Job Title: "] = fix_title(doc["title"])
        new_doc["Job Description: "] = fix_description(doc["desc"])
        new_doc["Qualifications: "] = fix_qualifications(doc["qual"])
        new_doc["Type: "] = fix_type(doc["type"])
        new_doc["How to Apply: "] = doc["how"]
        new_doc["Salary: "] = fix_salary(doc["salary"])
        new_doc["Contact: "] = fix_contact(doc["contact"])

        new_doc_list.append(new_doc)
        new_doc = {}

    return new_doc_list


def create_new_document(new_doc_list):
    new_document_list = []

    for doc in new_doc_list:
        new_document = Document()
        style = new_document.styles["Normal"]
        font = style.font
        font.name = "Arial"
        hstyle = new_document.styles["Heading 2"]
        hstyle.font.name = "Arial"
        hstyle.font.size = Pt(16)
        hstyle.font.color.rgb = RGBColor(0, 0, 0)

        h = new_document.add_heading("", 2)
        h.add_run(doc["Job Title: "] + ", " + doc["Submission Date: "]).bold = False

        for k, v in doc.items():
            p = new_document.add_paragraph()
            p.style = new_document.styles["Normal"]
            p.add_run(k).bold = True
            p.add_run(v)
            p.add_run()

        for paragraph in new_document.paragraphs:
            if paragraph.style.name.startswith("Heading 2"):
                for run in paragraph.runs:
                    run.font.name = "Arial"

        new_document_list.append(new_document)

    return new_document_list


def save_documents(new_document_list):
    for doc in new_document_list:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "  " in run.text:
                    run.text = run.text.replace("  ", " ")
                if run.text.startswith(" "):
                    run.text = run.text[1:]

        doc.save(join(current_dir, str(new_document_list.index(doc)) + ".docx"))


def combine_documents(documents):
    combined_doc = Document()
    heading = combined_doc.styles["Heading 2"]
    heading.font.name = "Arial"
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor(0, 0, 0)

    body = combined_doc.styles["Normal"]
    body.font.name = "Arial"

    for doc in documents:
        if isinstance(doc, int):
            doc_name = str(doc) + ".docx"
        else:
            doc_name = doc + ".docx"
        docf = Document(join(current_dir, doc_name))

        for element in docf.element.body:
            combined_doc.element.body.append(element)

    combined_doc.save(join(current_dir, "combined.docx"))


fix_document = load_fix_document("job-list.docx")
new_document = create_new_document(fix_document)
save_documents(new_document)
combine_documents(
    [
        0,
        1,
        2,
        3,
        4,
        5,
        6,
        7,
        8,
        9,
        10,
        11,
        12,
        13,
        14,
        15,
        16,
        17,
        18,
        19,
        20,
        21,
        22,
        23,
    ]
)

# new_document.save(join(current_dir, "test.docx"))
